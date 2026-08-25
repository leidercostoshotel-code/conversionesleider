#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
pdf_a_word.py
=============
Convierte los manuales de procedimientos "SPI Swissôtel Lima" (PDF escaneado,
sin capa de texto) a un Word (.docx) EDITABLE que replica el diseño original:
la tabla de cabecera (Área / Campo / Norma / N° / fechas) y el recuadro con
Objetivo / Responsables / Pasos a Seguir.

------------------------------------------------------------------------
INSTALACIÓN (una sola vez)
------------------------------------------------------------------------
Windows (con Chocolatey) / Mac (con Homebrew) / Linux (con apt):

    # Motor de OCR + idioma español
    apt install tesseract-ocr tesseract-ocr-spa poppler-utils      # Linux
    brew install tesseract tesseract-lang poppler                  # Mac
    choco install tesseract poppler                                 # Windows
                                                                     (en Windows,
                                                                     agrega también el
                                                                     paquete de idioma
                                                                     "spa" desde
                                                                     github.com/tesseract-ocr/tessdata)

    # Librerías de Python
    pip install python-docx pytesseract pdf2image pillow

------------------------------------------------------------------------
USO
------------------------------------------------------------------------
    python pdf_a_word.py archivo.pdf
    python pdf_a_word.py archivo.pdf -o salida.docx
    python pdf_a_word.py archivo.pdf --paginas 2-3        # solo esas páginas del PDF
    python pdf_a_word.py carpeta_con_pdfs/                # convierte todos los .pdf de la carpeta

------------------------------------------------------------------------
CÓMO FUNCIONA (y sus límites)
------------------------------------------------------------------------
1. Cada página del PDF se rasteriza a imagen (300 dpi).
2. Se le aplica OCR en español (Tesseract).
3. Se busca en el texto el patrón de cabecera fijo de estos manuales
   (SPI SWISSÔTEL LIMA / N° / Área / Campo / Norma / fechas). Si aparece,
   se reconstruye esa página como texto real: título en negrita, listas
   numeradas o con viñetas, subtítulos ("Recepción", "Puntos de Venta", etc.).
4. Si una página NO tiene ese patrón (formularios, organigramas, cuadros
   financieros de Excel, portadas), se inserta la página TAL CUAL como
   imagen, en vez de forzar un OCR que saldría mal. Así no se pierde
   información aunque no quede editable.
5. El OCR nunca es 100% exacto (fechas, siglas y nombres propios son los
   que más fallan). SIEMPRE revisa el .docx resultante antes de usarlo;
   piensa en este programa como quien te hace el primer borrador rápido,
   no como un reemplazo de la revisión final.

Este programa fue hecho a la medida del formato de estos manuales
("SPI SWISSÔTEL LIMA — Manual de Procedimientos"). Con otro tipo de PDF
probablemente no reconozca la cabecera y meta todo como imagen — lo cual
sigue siendo seguro (no pierde contenido), solo que no queda editable.
"""

import argparse
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

FONT = "Arial"
DPI = 300

# ------------------------------------------------------------------ #
# 1. PDF -> imágenes, imagen -> texto (OCR)
# ------------------------------------------------------------------ #

def pdf_to_images(pdf_path: Path, workdir: Path, dpi: int = DPI, page_range=None):
    """Rasteriza el PDF a una imagen PNG por página con pdftoppm (poppler)."""
    prefix = workdir / "page"
    cmd = ["pdftoppm", "-png", "-r", str(dpi)]
    if page_range:
        cmd += ["-f", str(page_range[0]), "-l", str(page_range[1])]
    cmd += [str(pdf_path), str(prefix)]
    subprocess.run(cmd, check=True, capture_output=True)
    images = sorted(workdir.glob("page-*.png"))
    if not images:
        # pdftoppm no rellena ceros si hay una sola página
        images = sorted(workdir.glob("page*.png"))
    return images


def ocr_page(image_path: Path) -> str:
    """OCR en español sobre una imagen de página. Devuelve el texto plano.

    Se usa --psm 4 ("una sola columna de texto de tamaños variables") en
    vez de --psm 6 ("bloque uniforme de texto): con --psm 6, Tesseract
    tiende a fusionar la tabla de cabecera (dos columnas) en una sola
    línea por fila y así se pierde por completo la segunda línea de esa
    tabla (donde está "Página : X/Y"), lo que hacía que la página cayera
    al modo imagen aunque el formato fuera reconocible.
    """
    result = subprocess.run(
        ["tesseract", str(image_path), "-", "-l", "spa", "--psm", "4"],
        check=True, capture_output=True, text=True,
    )
    return result.stdout


# ------------------------------------------------------------------ #
# 2. Reconocer la cabecera fija del manual
# ------------------------------------------------------------------ #

RIGHT_LABELS = [
    ("fecha_elaboracion", r"Fecha de elaboraci[oó0]n"),
    ("reemplaza", r"Reemplaza a la norma del"),
    ("referencia_corporativa", r"Referencia Corporativa"),
    ("elaborada_por", r"Elaborada por"),
    ("revisado_aprobado", r"Revisado y aprobado por"),
    ("autorizado_por", r"Autorizado por el Gerente General"),
]
LEFT_LABELS = [
    ("area", r"[AÁ]rea"),
    ("campo", r"Campo"),
    ("norma_titulo", r"Norma"),
]
ALL_LABELS = LEFT_LABELS + RIGHT_LABELS
LABEL_RE = re.compile("|".join(f"(?P<{k}>{p})" for k, p in ALL_LABELS))

NUMERO_RE = re.compile(r"\b(\d{2}\.\d{2}(?:\.\d{2}){0,2})\b")
PAGINA_RE = re.compile(r"P[aá]gina[^\d]{0,40}(\d+\s*/\s*\d+)")


def parse_header(text: str):
    """
    Busca el bloque de cabecera (primeras líneas, antes de 'OBJETIVO' /
    'PASOS A SEGUIR') y extrae sus campos. Devuelve (dict_o_None, resto_texto).
    """
    lines = text.splitlines()
    body_start_idx = None
    for i, ln in enumerate(lines):
        if re.search(r"OBJETIVO\s*:", ln, re.I) or re.search(r"PASOS A SEGUIR", ln, re.I):
            body_start_idx = i
            break

    if body_start_idx is None:
        # Página de continuación de un procedimiento de varias páginas: la
        # cabecera se repite en cada página, pero "OBJETIVO:"/"PASOS A
        # SEGUIR" solo aparecen en la primera. Antes, al no encontrar esos
        # marcadores, se descartaba TODO el cuerpo de la página (quedaba
        # solo la cabecera y el resto en blanco). En vez de eso, se ubica
        # el final real de la cabecera: la última línea con alguna etiqueta
        # fija, más alguna línea corta de continuación (valores envueltos
        # a varias líneas, como "Norma" o "Autorizado por...").
        last_label_idx = None
        for i, ln in enumerate(lines[:15]):
            if LABEL_RE.search(ln):
                last_label_idx = i
        if last_label_idx is None:
            body_start_idx = min(10, len(lines))
        else:
            idx = last_label_idx + 1
            while (
                idx < len(lines)
                and idx <= last_label_idx + 2
                and lines[idx].strip()
                and len(lines[idx].strip()) <= 50
                and not NUM_ITEM_RE.match(lines[idx])
            ):
                idx += 1
            body_start_idx = idx

    header_lines = lines[:body_start_idx]
    header_text = "\n".join(header_lines)

    numero_match = NUMERO_RE.search(header_text)
    pagina_match = PAGINA_RE.search(header_text)
    if not numero_match or not pagina_match:
        return None, text  # no se reconoce cabecera -> esta página va como imagen

    fields = {"numero": numero_match.group(1), "pagina": pagina_match.group(1).replace(" ", "")}
    last_key = None
    for ln in header_lines:
        matches = list(LABEL_RE.finditer(ln))
        if not matches:
            if last_key:
                fields[last_key] = (fields.get(last_key, "") + " " + ln.strip()).strip()
            continue
        for idx, m in enumerate(matches):
            key = m.lastgroup
            val_start = m.end()
            val_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(ln)
            value = ln[val_start:val_end]
            value = re.sub(r"^[\s:]+", "", value).strip()
            fields[key] = (fields.get(key, "") + " " + value).strip() if key in fields and key not in ("area", "campo") else value
            last_key = key

    body_text = "\n".join(lines[body_start_idx:])
    return fields, body_text


# ------------------------------------------------------------------ #
# 3. Reconocer el cuerpo: Objetivo / Responsables / Pasos a Seguir
# ------------------------------------------------------------------ #

NUM_ITEM_RE = re.compile(r"^\s*(\d{1,2})[.\)]\s+(.*)")
LETTER_ITEM_RE = re.compile(r"^\s*([a-h])\)\s+(.*)")
BULLET_ITEM_RE = re.compile(r"^\s*[•·\-\*]\s+(.*)")

# Ruido típico de capturas de un visor web (fecha/hora, migas de pan,
# barra de usuario, pie con URL y contador de página) que se mete en el
# OCR cuando el PDF es en realidad una captura de pantalla de un visor.
NOISE_SUB_PATTERNS = [
    r"https?\s*:\s*//\S*[\w/=?.]*",
    r"\d{1,2}/\d{1,2}/\d{2,4},?\s*\d{1,2}:\d{2}\s*[ap]\.?\s*m\.?\.?",
    r"\b\d{1,2}\s*/\s*\d{1,2}\b(?=\s*$)",
]
NOISE_LINE_PATTERNS = [
    r"Intranet corporativa",
    r"Cerrar sesi[oó]n",
    r"—\s*Intranet\s*$",
    r"^\s*[←<]\s*\w+\s*[-–—]?\s*\d{0,3}%?\s*[-+]?\s*$",
    r"^\s*P[aá]g\.?\s*\d+\s*$",
]


def strip_noise(text: str) -> str:
    for pat in NOISE_SUB_PATTERNS:
        text = re.sub(pat, " ", text)
    lines = text.splitlines()
    lines = [ln for ln in lines if not any(re.search(p, ln, re.I) for p in NOISE_LINE_PATTERNS)]
    return "\n".join(lines)


def is_heading_line(line: str) -> bool:
    """Título de sección corto, con mayúscula inicial y terminado en ':'
    (p. ej. "RESPONSABLES:", "BASE LEGAL:", "CARACTERISTICAS:")."""
    line = line.strip()
    return bool(line) and len(line) <= 45 and line.endswith(":") and line[:1].isupper()


def split_section(text, start_pat, end_pats):
    """Devuelve (prefijo_antes_de_la_marca, valor_de_la_sección, resto).
    `prefijo` es lo que había ANTES de encontrar `start_pat` -- normalmente
    vacío, salvo cuando hay una sección intermedia no prevista (p. ej.
    "BASE LEGAL:" entre "OBJETIVO:" y "RESPONSABLES:") que de otro modo se
    perdería en silencio. Si no se encuentra `start_pat`, devuelve
    (None, None, text) sin consumir nada."""
    m = re.search(start_pat, text, re.I)
    if not m:
        return None, None, text
    prefix = text[:m.start()]
    rest = text[m.end():]
    end_pos = len(rest)
    for ep in end_pats:
        em = re.search(ep, rest, re.I)
        if em and em.start() < end_pos:
            end_pos = em.start()
    # Además de los patrones explícitos (que solo cubren los nombres de
    # sección esperados), cualquier línea que en sí misma parezca un
    # título de sección marca el final de este bloque. Esto generaliza a
    # secciones no previstas (p. ej. "BASE LEGAL:", "CARACTERISTICAS:")
    # que de otro modo quedarían arrastradas dentro de "OBJETIVO" o
    # "RESPONSABLES" hasta la siguiente sección conocida.
    pos = 0
    for line in rest.splitlines(keepends=True):
        if pos >= end_pos:
            break
        if pos > 0 and is_heading_line(line):
            end_pos = pos
            break
        pos += len(line)
    return prefix, rest[:end_pos].strip(), rest[end_pos:]


def parse_body(body_text: str):
    """
    Devuelve una lista de bloques tipados:
      {'type': 'objetivo'|'responsables'|'subheading'|'numbered'|'bullet'|'letter'|'paragraph', ...}
    """
    blocks = []
    remaining = body_text

    pre_objetivo, objetivo, remaining = split_section(
        remaining, r"OBJETIVO\s*:", [r"RESPONSABLES\s*:", r"PASOS A SEGUIR", r"POL[IÍ]TICAS"]
    )
    if pre_objetivo and pre_objetivo.strip():
        blocks.extend(parse_steps(pre_objetivo))
    if objetivo:
        blocks.append({"type": "objetivo", "text": clean_paragraph(objetivo)})

    pre_responsables, responsables, remaining = split_section(
        remaining, r"RESPONSABLES\s*:", [r"PASOS A SEGUIR", r"POL[IÍ]TICAS"]
    )
    if pre_responsables and pre_responsables.strip():
        # Sección intermedia no prevista (p. ej. "BASE LEGAL:") entre
        # "OBJETIVO:" y "RESPONSABLES:": se reconstruye con el parser
        # genérico en vez de perderla.
        blocks.extend(parse_steps(pre_responsables))
    if responsables:
        names = [ln.strip(" .") for ln in responsables.splitlines() if ln.strip()]
        blocks.append({"type": "responsables", "names": names})

    # Lo que sigue (o todo el texto, si esta página es una continuación
    # sin encabezados) se interpreta como la lista de pasos.
    steps_text = remaining
    m = re.search(r"PASOS A SEGUIR\s*:?", steps_text, re.I)
    if m:
        steps_text = steps_text[m.end():]

    blocks.extend(parse_steps(steps_text))
    return blocks


def clean_paragraph(text):
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    return " ".join(lines)


def parse_steps(text):
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    blocks = []
    current = None
    i = 0
    while i < len(lines):
        line = lines[i]

        m = NUM_ITEM_RE.match(line)
        if m:
            if current:
                blocks.append(current)
            current = {"type": "numbered", "n": m.group(1), "text": m.group(2).strip()}
            i += 1
            continue

        m = LETTER_ITEM_RE.match(line)
        if m:
            if current:
                blocks.append(current)
            current = {"type": "letter", "letter": m.group(1), "text": m.group(2).strip()}
            i += 1
            continue

        m = BULLET_ITEM_RE.match(line)
        if m:
            if current:
                blocks.append(current)
            current = {"type": "bullet", "text": m.group(1).strip()}
            i += 1
            continue

        # Un título de sección corto, en mayúscula inicial y terminado en
        # ":" (p. ej. "CONTROLES:", "REMUNERACIONES:", "PAGOS:") es en sí
        # mismo una señal inequívoca de subtítulo en este tipo de manuales
        # -- no hace falta que además reinicie una lista en "1.".
        if is_heading_line(line):
            if current:
                blocks.append(current)
                current = None
            blocks.append({"type": "subheading", "text": line})
            i += 1
            continue

        looks_like_heading = (
            len(line) <= 45
            and not line.endswith((".", ",", ";"))
            and line[:1].isupper()
        )
        # Un subtítulo sin ":" casi siempre viene seguido de un ítem "1."
        # (la lista de esa sub-sección vuelve a empezar). Usamos esa
        # señal aunque haya un ítem numerado abierto en ese momento.
        next_starts_list_at_1 = False
        if i + 1 < len(lines):
            nm = NUM_ITEM_RE.match(lines[i + 1])
            next_starts_list_at_1 = bool(nm and nm.group(1) == "1")

        if looks_like_heading and (current is None or next_starts_list_at_1):
            if current:
                blocks.append(current)
                current = None
            blocks.append({"type": "subheading", "text": line})
            i += 1
            continue

        if current:
            current["text"] = (current["text"] + " " + line).strip()
        else:
            blocks.append({"type": "paragraph", "text": line})
        i += 1

    if current:
        blocks.append(current)
    return blocks


# ------------------------------------------------------------------ #
# 4. Construcción del .docx (mismo diseño que el original)
# ------------------------------------------------------------------ #

def set_cell_border(cell, size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tc_pr.append(borders)


def add_run(paragraph, text, bold=False, size=10):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.space_before = Pt(0)
    return run


def add_header_table(doc, f):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Twips(6800), Twips(3100)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
            set_cell_border(cell)

    c = table.cell(0, 0)
    c.paragraphs[0].text = ""
    add_run(c.paragraphs[0], "SPI SWISSÔTEL LIMA", bold=True, size=11)
    p2 = c.add_paragraph()
    add_run(p2, "Manual de Procedimientos", size=10)

    c = table.cell(0, 1)
    add_run(c.paragraphs[0], f"N°  : {f.get('numero', '')}", bold=True, size=10)
    p2 = c.add_paragraph()
    add_run(p2, f"Página: {f.get('pagina', '')}", size=10)

    c = table.cell(1, 0)
    add_run(c.paragraphs[0], f"Area  : {f.get('area', '')}", size=9)
    for key, label in (("campo", "Campo"), ("norma_titulo", "Norma")):
        p = c.add_paragraph()
        add_run(p, f"{label} : {f.get(key, '')}", size=9)

    c = table.cell(1, 1)
    right_fields = [
        ("fecha_elaboracion", "Fecha de elaboración"),
        ("reemplaza", "Reemplaza a la norma del"),
        ("elaborada_por", "Elaborada por"),
        ("referencia_corporativa", "Referencia Corporativa SL"),
        ("revisado_aprobado", "Revisado y aprobado por"),
    ]
    first = True
    for key, label in right_fields:
        if not f.get(key):
            continue
        p = c.paragraphs[0] if first else c.add_paragraph()
        add_run(p, f"{label} : {f[key]}", size=9)
        first = False
    p = c.paragraphs[0] if first else c.add_paragraph()
    add_run(p, "Autorizado por el Gerente General :", size=9)
    return table


def add_content_box(doc, blocks):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Twips(9900)
    set_cell_border(cell)
    cell.paragraphs[0].text = ""
    first = True

    def new_para():
        nonlocal first
        if first:
            first = False
            return cell.paragraphs[0]
        return cell.add_paragraph()

    for b in blocks:
        t = b["type"]
        if t == "objetivo":
            add_run(new_para(), "OBJETIVO:", bold=True)
            add_run(new_para(), b["text"])
        elif t == "responsables":
            add_run(new_para(), "RESPONSABLES:", bold=True)
            for name in b["names"]:
                add_run(new_para(), name)
        elif t == "subheading":
            add_run(new_para(), b["text"], bold=True)
        elif t == "numbered":
            p = new_para()
            p.paragraph_format.left_indent = Twips(360)
            add_run(p, f"{b['n']}.  {b['text']}")
        elif t == "letter":
            p = new_para()
            p.paragraph_format.left_indent = Twips(500)
            add_run(p, f"{b['letter']})  {b['text']}")
        elif t == "bullet":
            p = new_para()
            p.paragraph_format.left_indent = Twips(360)
            add_run(p, f"•  {b['text']}")
        else:  # paragraph
            add_run(new_para(), b["text"])
    return table


def add_full_page_image(doc, image_path):
    section = doc.sections[-1]
    usable_w = section.page_width - section.left_margin - section.right_margin
    doc.add_picture(str(image_path), width=usable_w)


def add_page_break(doc):
    doc.add_page_break()


# ------------------------------------------------------------------ #
# 5. Orquestación
# ------------------------------------------------------------------ #

def parse_page_range(spec, total_pages):
    if not spec:
        return None
    if "-" in spec:
        a, b = spec.split("-", 1)
        return (int(a), int(b))
    n = int(spec)
    return (n, n)


def convert(pdf_path: Path, out_path: Path, page_range_spec=None, keep_tmp=False, verbose=True):
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        page_range = parse_page_range(page_range_spec, None)
        images = pdf_to_images(pdf_path, tmp, page_range=page_range)
        if not images:
            print(f"  ! No se pudo rasterizar {pdf_path.name}", file=sys.stderr)
            return

        doc = Document()
        section = doc.sections[0]
        section.page_width, section.page_height = Twips(11906), Twips(16838)  # A4
        section.top_margin = section.bottom_margin = Twips(700)
        section.left_margin = section.right_margin = Twips(700)

        for i, img_path in enumerate(images):
            if verbose:
                print(f"  · OCR página {i + 1}/{len(images)}...")
            text = ocr_page(img_path)
            text = strip_noise(text)
            fields, body_text = parse_header(text)

            if i > 0:
                add_page_break(doc)

            if fields is None:
                # No se reconoció la cabecera -> se inserta la página tal cual
                add_full_page_image(doc, img_path)
                continue

            add_header_table(doc, fields)
            doc.add_paragraph()  # espaciador
            blocks = parse_body(body_text)
            add_content_box(doc, blocks)

        doc.save(out_path)
        if verbose:
            print(f"  -> {out_path}")


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("entrada", help="Archivo .pdf o carpeta con varios .pdf")
    ap.add_argument("-o", "--salida", help="Ruta del .docx de salida (solo si 'entrada' es un archivo)")
    ap.add_argument("--paginas", help="Rango de páginas del PDF a convertir, ej: 2-3")
    args = ap.parse_args()

    entrada = Path(args.entrada)
    if entrada.is_dir():
        pdfs = sorted(entrada.glob("*.pdf"))
        if not pdfs:
            print("No hay archivos .pdf en esa carpeta.")
            return
        for pdf in pdfs:
            print(f"Procesando {pdf.name}...")
            out = pdf.with_suffix(".docx")
            convert(pdf, out, args.paginas)
    else:
        out = Path(args.salida) if args.salida else entrada.with_suffix(".docx")
        print(f"Procesando {entrada.name}...")
        convert(entrada, out, args.paginas)

    print("Listo. Revisa el/los .docx generados antes de usarlos (el OCR puede tener errores).")


if __name__ == "__main__":
    main()
